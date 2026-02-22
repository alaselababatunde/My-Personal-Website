import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.name = 'Arial'
    run.font.color.rgb = docx.shared.RGBColor(0, 51, 102)

def set_normal_font(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

doc = docx.Document()
set_normal_font(doc)

# Header
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("ALASELA BABATUNDE\n")
run.bold = True
run.font.size = Pt(20)
run.font.name = 'Arial'
run.font.color.rgb = docx.shared.RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("AI Engineer & Founder\n")
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Arial'

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.add_run("Lagos, Nigeria | 09154357208 | alaselababatunde10@gmail.com\n")
contact.add_run("LinkedIn: linkedin.com/in/alasela-babatunde-769524333 | GitHub: github.com/alaselababatunde")

doc.add_paragraph() # spacing

# Summary
add_heading(doc, "PROFESSIONAL SUMMARY", 1)
summary = doc.add_paragraph("AI Engineer & Founder (CuraAi, Co) specializing in emotionally intelligent AI systems, RAG pipelines, and modern web applications. "
                            "Passionate about building AI solutions that are ethical, accessible, and impactful — proving that age is just a number when dedication meets innovation. "
                            "Proven hackathon winner and AI-Python educator.")

# Skills
add_heading(doc, "TECHNICAL SKILLS", 1)
skills = doc.add_paragraph()
skills.add_run("Languages: ").bold = True
skills.add_run("Python, HTML, C++\n")
skills.add_run("Frameworks & AI: ").bold = True
skills.add_run("Flask, Langchain, Ollama, Hugging Face, OpenAI SDK, OpenRouter SDK\n")
skills.add_run("Tools: ").bold = True
skills.add_run("Git, Docker, Vercel, Hugging Face Spaces\n")
skills.add_run("Domains: ").bold = True
skills.add_run("AI for Mental Health, AI for Agriculture, Generative AI, Web Development, RAG Systems")

# Experience
add_heading(doc, "PROFESSIONAL EXPERIENCE", 1)

experiences = [
    ("Founder & AI Engineer", "Cura AI, Co.", "2025 – Present", "Leading development of emotionally intelligent AI virtual companions and RAG systems for mental health support."),
    ("Founder, Solution Architect & AI Engineer", "Alash Studios", "2023 – Present", "Founded and lead a Tech Studio focused on AI innovation. Designing and deploying custom AI solutions and optimization tools."),
    ("Lead AI Engineer (Hackathon Winner)", "DevAssist Team", "Oct 2025", "Won Remostart Hackathon by building an AI-powered IDE with code debugging and auto-documentation features."),
    ("AI Engineer (FutureStack GenAI)", "Team Astra", "2025", "Developed AI solutions for agricultural optimization and crop prediction during the FutureStack GenAI Hackathon."),
    ("AI Developer", "Tech Disciples AI", "2025", "Created custom AI assistants to support community engagement and information access for the Tech Disciples church community."),
    ("Python & Prompt Engineering Tutor", "CLCC Digital Bootcamp", "2025", "Trained individuals on Python programming and prompt engineering."),
    ("Project Supervisor", "Luca Ador", "Aug 2024 – Dec 2024", "Overseeing multiple technology projects, ensuring timely completion and technical quality across teams."),
    ("Python & AI Instructor", "American Corner", "2023 – 2024", "Trained individuals and teams on Python programming and AI integration strategies."),
    ("Python Developer", "A.S.I.A Project", "2023 – 2024", "Built a Python-based virtual assistant (A.S.I.A) using speech recognition, TTS, and real-time animations.")
]

for role, company, period, desc in experiences:
    p = doc.add_paragraph()
    p.add_run(role).bold = True
    p.add_run(f" | {company}")
    run = p.add_run(f" ({period})")
    run.font.italic = True
    
    desc_p = doc.add_paragraph(desc)
    desc_p.style = 'List Bullet'

# Projects
add_heading(doc, "SELECTED PROJECTS & A.I. SYSTEMS", 1)

projects = [
    ("UBA AI Support", "Banking", "A production-grade customer support AI system for UBA Bank featuring a RAG pipeline and session-based memory."),
    ("Tesco AI Support", "E-commerce", "Intelligent customer support RAG system for Tesco, optimizing customer interactions with AI-driven responses."),
    ("DSTV AI Support", "Entertainment", "A specialized AI customer support system for DStv, leveraging RAG technology to provide instant, accurate assistance."),
    ("Domino's Pizza AI Support", "Food Tech", "Intelligent customer support AI system for Domino's Pizza, providing instant help for ordering and inquiries."),
    ("Theo AI", "Theology / Christian", "A modern Christian AI assistant designed to provide scriptural guidance and spiritual support through an intelligent RAG-powered interface."),
    ("Premium Static Websites", "Frontend / UI/UX", "Developed high-end, smooth vanilla JS architectures for Echo Microfinance Bank, Romax Properties, and Sales Yards Nigeria.")
]

for title, tags, p_desc in projects:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run(title).bold = True
    p.add_run(f" ({tags}): ").font.italic = True
    p.add_run(p_desc)

# Education & Certifications
add_heading(doc, "EDUCATION & CERTIFICATIONS", 1)

education = [
    ("Stars College", "High School Diploma (Graduating 2026)"),
    ("Hugging Face", "The LLM Course"),
    ("Google / Coursera", "Google AI Essentials"),
    ("Coursera", "Rasa Chatbot Development, Python (GUI & OOP)"),
    ("EdX", "Python Programming (Beginner), AI With Python"),
    ("Stellar Academy", "Cybersecurity")
]

for inst, cert in education:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run(f"{inst}: ").bold = True
    p.add_run(cert)

# Achievements
add_heading(doc, "AWARDS & ACHIEVEMENTS", 1)

achievements = [
    "Hackathon Winner — Remostart Hackathon (DevAssist Team) — October 2025",
    "Top 20 Finalist — Nigeria AI Digital Art Competition — 2025",
    "Tech Titan of Cohort 3.0 — Stellar Academy",
    "Laptop & Scholarship Award — Boycode 3.0 for Excellence in Tech"
]

for ach in achievements:
    p = doc.add_paragraph(ach)
    p.style = 'List Bullet'

doc.save("/home/alash-studios/My-Personal-Website/Alasela_Babatunde_Updated_Resume.docx")
print("Resume generated successfully.")
